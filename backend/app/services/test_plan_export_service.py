"""Test plan export service — generates PDF and DOCX reports."""

import io
import logging
from datetime import datetime
from typing import Any, Dict, List, Optional

from app.models.test_plan import TestPlan

logger = logging.getLogger(__name__)

# ============================================================
# SECTION DEFINITIONS
# ============================================================

_SECTIONS = [
    ("description", "1. Description"),
    ("objective", "2. Objective"),
    ("in_scope", "3. In Scope"),
    ("out_of_scope", "4. Out of Scope"),
    ("entry_criteria", "5. Entry Criteria"),
    ("exit_criteria", "6. Exit Criteria"),
    ("approach", "7. Test Approach"),
    ("assumptions", "8. Assumptions"),
    ("constraints", "9. Constraints"),
    ("stakeholders", "10. Stakeholders & Responsibilities"),
    ("communication", "11. Communication Plan"),
]

_META_ROWS = [
    ("Environment", "environment"),
    ("Scope Type", "scope_type"),
]


class TestPlanExportService:

    # ============================================================
    # PDF
    # ============================================================

    def export_pdf(self, plan: TestPlan) -> bytes:
        """Generate a professional PDF report for the test plan."""
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import mm
            from reportlab.lib import colors
            from reportlab.platypus import (
                SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
                HRFlowable, KeepTogether,
            )
            from reportlab.lib.enums import TA_LEFT, TA_CENTER

        except ImportError as exc:
            logger.error(f"[EXPORT] reportlab not installed: {exc}")
            raise RuntimeError(
                "PDF export requires 'reportlab'. "
                "Install it with: pip install reportlab"
            ) from exc

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            leftMargin=20 * mm,
            rightMargin=20 * mm,
            topMargin=20 * mm,
            bottomMargin=20 * mm,
        )

        styles = getSampleStyleSheet()
        brand_blue = colors.HexColor("#1a6fc4")
        dark = colors.HexColor("#1a1a2e")

        h1 = ParagraphStyle(
            "H1",
            parent=styles["Heading1"],
            fontSize=22,
            textColor=dark,
            spaceAfter=4,
        )
        h2 = ParagraphStyle(
            "H2",
            parent=styles["Heading2"],
            fontSize=13,
            textColor=brand_blue,
            spaceBefore=14,
            spaceAfter=4,
            borderPad=2,
        )
        body = ParagraphStyle(
            "Body",
            parent=styles["Normal"],
            fontSize=10,
            leading=15,
            spaceAfter=6,
        )
        meta_label = ParagraphStyle(
            "MetaLabel",
            parent=styles["Normal"],
            fontSize=9,
            textColor=colors.grey,
            spaceAfter=2,
        )
        status_style = ParagraphStyle(
            "Status",
            parent=styles["Normal"],
            fontSize=9,
            textColor=brand_blue,
            spaceAfter=6,
        )

        story = []

        # ── Header ──────────────────────────────────────────────
        story.append(Paragraph("TESTFORGE", ParagraphStyle(
            "Brand", parent=styles["Normal"],
            fontSize=10, textColor=brand_blue, spaceAfter=2,
        )))
        story.append(Paragraph(plan.title, h1))
        story.append(Paragraph(
            f"Status: <b>{plan.status.replace('_', ' ').title()}</b>   |   "
            f"Generated: {datetime.utcnow().strftime('%Y-%m-%d')}",
            status_style,
        ))
        story.append(HRFlowable(width="100%", thickness=1, color=brand_blue, spaceAfter=10))

        # ── Metadata table ───────────────────────────────────────
        meta_data = []
        for label, attr in _META_ROWS:
            value = getattr(plan, attr, None)
            if value:
                if isinstance(value, list):
                    value = ", ".join(value)
                meta_data.append([label, value])
        if plan.start_date:
            meta_data.append(["Start Date", str(plan.start_date)])
        if plan.end_date:
            meta_data.append(["End Date", str(plan.end_date)])

        if meta_data:
            table = Table(
                meta_data,
                colWidths=[45 * mm, 120 * mm],
                hAlign="LEFT",
            )
            table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#f0f4ff")),
                ("TEXTCOLOR", (0, 0), (0, -1), colors.HexColor("#555555")),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
                ("ROWBACKGROUNDS", (0, 0), (-1, -1), [colors.white, colors.HexColor("#fafbff")]),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#dddddd")),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]))
            story.append(table)
            story.append(Spacer(1, 10))

        # ── Risk Analysis Section ──────────────────────────────────
        if plan.risk_analysis:
            story.append(Paragraph("Risk Analysis", h2))
            story.append(HRFlowable(width="100%", thickness=0.5, color=colors.red, spaceAfter=4))
            
            # Distribution
            dist = plan.risk_analysis.get('distribution', {})
            if dist:
                dist_data = [
                    ["Critical", str(dist.get('critical', 0)), "🔴"],
                    ["High", str(dist.get('high', 0)), "🟠"],
                    ["Medium", str(dist.get('medium', 0)), "🟡"],
                    ["Low", str(dist.get('low', 0)), "🟢"],
                    ["Total", str(dist.get('total', 0)), ""],
                ]
                dist_table = Table(dist_data, colWidths=[60*mm, 40*mm, 20*mm])
                dist_table.setStyle(TableStyle([
                    ("FONTSIZE", (0, 0), (-1, -1), 9),
                    ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#f0f4ff")),
                    ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#dddddd")),
                ]))
                story.append(dist_table)
                story.append(Spacer(1, 6))
            
            # Risk Mapping Table (Critical & High only)
            mapping = plan.risk_analysis.get('mapping_table', [])
            critical_high = [r for r in mapping if r.get('risk_level') in ('critical', 'high')]
            if critical_high:
                mapping_data = [["Key", "Title", "Score", "Mitigation"]]
                for r in critical_high[:10]:
                    mapping_data.append([
                        r.get('issue_key', '?'),
                        r.get('title', '')[:40],
                        str(r.get('risk_score', 0)),
                        r.get('mitigation', '')[:80],
                    ])
                mapping_table = Table(mapping_data, colWidths=[25*mm, 50*mm, 15*mm, 70*mm])
                mapping_table.setStyle(TableStyle([
                    ("FONTSIZE", (0, 0), (-1, -1), 8),
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#fef2f2")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.red),
                    ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#dddddd")),
                ]))
                story.append(mapping_table)
        # ── Sections ─────────────────────────────────────────────
        for attr, label in _SECTIONS:
            value = getattr(plan, attr, None)
            if not value:
                continue
            content = [
                Paragraph(label, h2),
                HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#dddddd"), spaceAfter=4),
            ]
            for line in str(value).split("\n"):
                line = line.strip()
                if line:
                    content.append(Paragraph(line, body))
            story.append(KeepTogether(content))

        # ── Footer ───────────────────────────────────────────────
        story.append(Spacer(1, 16))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.lightgrey, spaceAfter=4))
        story.append(Paragraph(
            "Generated by TestForge AI — Intelligent Test Automation",
            ParagraphStyle("Footer", parent=styles["Normal"],
                           fontSize=8, textColor=colors.grey, alignment=TA_CENTER),
        ))

        doc.build(story)
        return buffer.getvalue()

    # ============================================================
    # DOCX
    # ============================================================

    def export_docx(self, plan: TestPlan) -> bytes:
        """Generate a professional DOCX report for the test plan."""
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement

        except ImportError as exc:
            logger.error(f"[EXPORT] python-docx not installed: {exc}")
            raise RuntimeError(
                "DOCX export requires 'python-docx'. "
                "Install it with: pip install python-docx"
            ) from exc

        doc = Document()

        # ── Page margins ─────────────────────────────────────────
        for section in doc.sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

        brand_blue = RGBColor(0x1A, 0x6F, 0xC4)
        dark = RGBColor(0x1A, 0x1A, 0x2E)

        # ── Brand header ─────────────────────────────────────────
        brand_p = doc.add_paragraph()
        brand_run = brand_p.add_run("TESTFORGE")
        brand_run.font.color.rgb = brand_blue
        brand_run.font.size = Pt(9)
        brand_run.font.bold = True

        # ── Title ────────────────────────────────────────────────
        title_p = doc.add_heading(plan.title, level=1)
        title_p.runs[0].font.color.rgb = dark

        # ── Status line ──────────────────────────────────────────
        status_p = doc.add_paragraph()
        status_run = status_p.add_run(
            f"Status: {plan.status.replace('_', ' ').title()}   |   "
            f"Generated: {datetime.utcnow().strftime('%Y-%m-%d')}"
        )
        status_run.font.color.rgb = brand_blue
        status_run.font.size = Pt(9)

        self._add_hr(doc)

        # ── Metadata table ───────────────────────────────────────
        meta_rows = []
        for label, attr in _META_ROWS:
            value = getattr(plan, attr, None)
            if value:
                if isinstance(value, list):
                    value = ", ".join(value)
                meta_rows.append((label, str(value)))
        if plan.start_date:
            meta_rows.append(("Start Date", str(plan.start_date)))
        if plan.end_date:
            meta_rows.append(("End Date", str(plan.end_date)))

        if meta_rows:
            table = doc.add_table(rows=len(meta_rows), cols=2)
            table.style = "Table Grid"
            for i, (label, value) in enumerate(meta_rows):
                cell_label = table.cell(i, 0)
                cell_value = table.cell(i, 1)
                cell_label.text = label
                cell_value.text = value
                # Style label cell
                for run in cell_label.paragraphs[0].runs:
                    run.font.bold = True
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
                for run in cell_value.paragraphs[0].runs:
                    run.font.size = Pt(9)
            doc.add_paragraph()

        # ── Content sections ─────────────────────────────────────
        for attr, label in _SECTIONS:
            value = getattr(plan, attr, None)
            if not value:
                continue

            h = doc.add_heading(label, level=2)
            h.runs[0].font.color.rgb = brand_blue

            for line in str(value).split("\n"):
                line = line.strip()
                if line:
                    p = doc.add_paragraph(line)
                    p.style.font.size = Pt(10)

        # ── Footer ───────────────────────────────────────────────
        self._add_hr(doc)
        footer_p = doc.add_paragraph("Generated by TestForge AI — Intelligent Test Automation")
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_p.runs[0].font.size = Pt(8)
        footer_p.runs[0].font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    # ============================================================
    # HELPERS
    # ============================================================

    @staticmethod
    def _add_hr(doc) -> None:
        """Add a horizontal rule paragraph to a docx document."""
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "AAAAAA")
        pBdr.append(bottom)
        pPr.append(pBdr)


# ============================================================
# SUITE EXECUTION REPORT
# ============================================================

class SuiteReportExportService:
    """Generates a detailed PDF execution report for a test suite run."""

    # ── status helpers ──────────────────────────────────────────
    _STATUS_COLOR = {
        "passed":  "#16a34a",
        "failed":  "#dc2626",
        "error":   "#ea580c",
        "skipped": "#6b7280",
    }
    _STATUS_BG = {
        "passed":  "#f0fdf4",
        "failed":  "#fef2f2",
        "error":   "#fff7ed",
        "skipped": "#f9fafb",
    }
    _STATUS_ICON = {"passed": "PASSED", "failed": "FAILED", "error": "ERROR", "skipped": "SKIPPED"}

    def export_pdf(
        self,
        suite_name: str,
        summary: Optional[Dict[str, Any]],
        entries: List[Dict[str, Any]],
        run_details: Dict[str, Any],
    ) -> bytes:
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import mm
            from reportlab.lib import colors
            from reportlab.platypus import (
                SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
                HRFlowable, KeepTogether,
            )
            from reportlab.lib.enums import TA_CENTER, TA_LEFT
        except ImportError as exc:
            raise RuntimeError(
                "PDF export requires 'reportlab'. Install with: pip install reportlab"
            ) from exc

        W = 170 * mm

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer, pagesize=A4,
            leftMargin=20 * mm, rightMargin=20 * mm,
            topMargin=20 * mm, bottomMargin=20 * mm,
        )

        styles = getSampleStyleSheet()
        BLUE  = colors.HexColor("#1a6fc4")
        DARK  = colors.HexColor("#1a1a2e")
        GREY  = colors.HexColor("#6b7280")
        LGREY = colors.HexColor("#e5e7eb")
        FLGREY = colors.HexColor("#f9fafb")
        GREEN = colors.HexColor("#16a34a")
        RED   = colors.HexColor("#dc2626")

        def _style(name, **kw):
            return ParagraphStyle(name, parent=styles["Normal"], **kw)

        S = {
            "brand":  _style("brand",  fontSize=9,  textColor=BLUE, spaceAfter=2),
            "h1":     _style("h1",     fontSize=22, textColor=DARK, spaceAfter=4, fontName="Helvetica-Bold"),
            "sname":  _style("sname",  fontSize=13, textColor=DARK, spaceAfter=4),
            "small":  _style("small",  fontSize=8,  textColor=GREY, spaceAfter=2),
            "footer": _style("footer", fontSize=8,  textColor=GREY, alignment=TA_CENTER),
        }

        def _hr(color=LGREY, thick=0.5, after=4):
            return HRFlowable(width="100%", thickness=thick, color=color, spaceAfter=after)

        story = []

        # ── Header ────────────────────────────────────────────
        story.append(Paragraph("TESTFORGE", S["brand"]))
        story.append(Paragraph("Suite Execution Report", S["h1"]))
        story.append(Paragraph(suite_name, S["sname"]))
        story.append(Paragraph(
            f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}",
            S["small"],
        ))
        story.append(_hr(BLUE, 1.5, 10))

        # ── Summary table ─────────────────────────────────────
        if summary:
            passed  = summary.get("passed", 0)
            failed  = summary.get("failed", 0)
            skipped = summary.get("skipped", 0)
            total   = passed + failed + skipped
            dur     = summary.get("duration", 0)

            _lbl = _style("sl", fontSize=9, textColor=GREY, fontName="Helvetica-Bold")
            sum_data = [
                [Paragraph("Total",    _lbl), Paragraph(str(total),    _style("sv0", fontSize=9, textColor=DARK))],
                [Paragraph("Passed",   _lbl), Paragraph(str(passed),   _style("sv1", fontSize=9, textColor=GREEN))],
                [Paragraph("Failed",   _lbl), Paragraph(str(failed),   _style("sv2", fontSize=9, textColor=RED))],
                [Paragraph("Skipped",  _lbl), Paragraph(str(skipped),  _style("sv3", fontSize=9, textColor=DARK))],
                [Paragraph("Duration", _lbl), Paragraph(f"{dur:.1f}s", _style("sv4", fontSize=9, textColor=DARK))],
            ]
            sum_tbl = Table(sum_data, colWidths=[55 * mm, 55 * mm], hAlign="LEFT")
            sum_tbl.setStyle(TableStyle([
                ("BACKGROUND",    (0, 0), (0, -1), FLGREY),
                ("GRID",          (0, 0), (-1, -1), 0.25, LGREY),
                ("LEFTPADDING",   (0, 0), (-1, -1), 8),
                ("RIGHTPADDING",  (0, 0), (-1, -1), 8),
                ("TOPPADDING",    (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ("VALIGN",        (0, 0), (-1, -1), "MIDDLE"),
            ]))
            story.append(sum_tbl)
            story.append(Spacer(1, 14))

        # ── Per test case ─────────────────────────────────────
        for entry in entries:
            run_id  = entry.get("tc_result_id") or entry.get("run_id")
            tc_code = entry.get("tc_code", "?")
            title   = entry.get("title", "Untitled")
            status  = (entry.get("status") or "unknown").lower()

            detail        = (run_details.get(run_id) or {}) if run_id else {}
            tc_result     = detail.get("tc_result") or {}
            exec_info     = detail.get("execution") or {}
            exec_steps    = detail.get("steps") or []
            justification = tc_result.get("justification") or ""

            browser  = exec_info.get("browser") or ""
            duration = tc_result.get("duration")

            if status == "passed":
                icon_char  = "✓"
                icon_color = GREEN
            elif status in ("failed", "error"):
                icon_char  = "✗"
                icon_color = RED
            else:
                icon_char  = "—"
                icon_color = GREY

            block = []
            block.append(Spacer(1, 8))

            # ── TC heading ────────────────────────────────────
            hdr_cells = [[
                Paragraph(
                    f"<b>{icon_char}</b>",
                    _style(f"ico_{tc_code}", fontSize=11, textColor=icon_color,
                           fontName="Helvetica-Bold"),
                ),
                Paragraph(
                    f"<b>{tc_code}</b> — {title}",
                    _style(f"tch_{tc_code}", fontSize=11, textColor=DARK,
                           fontName="Helvetica-Bold"),
                ),
            ]]
            hdr_tbl = Table(hdr_cells, colWidths=[7 * mm, W - 7 * mm], hAlign="LEFT")
            hdr_tbl.setStyle(TableStyle([
                ("LEFTPADDING",   (0, 0), (-1, -1), 0),
                ("RIGHTPADDING",  (0, 0), (-1, -1), 0),
                ("TOPPADDING",    (0, 0), (-1, -1), 2),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ("VALIGN",        (0, 0), (-1, -1), "MIDDLE"),
            ]))
            block.append(hdr_tbl)

            # ── Meta table ────────────────────────────────────
            _ml = _style("ml", fontSize=8, textColor=GREY)
            _mv = _style("mv", fontSize=8, textColor=DARK)

            meta_rows = [
                [Paragraph("Status",  _ml), Paragraph(status.upper(), _mv)],
                [Paragraph("Browser", _ml), Paragraph(browser or "—", _mv)],
            ]
            if duration is not None:
                meta_rows.append([Paragraph("Duration", _ml), Paragraph(f"{duration:.1f}s", _mv)])
            if justification:
                meta_rows.append([Paragraph("Result", _ml), Paragraph(justification[:300], _mv)])

            meta_tbl = Table(meta_rows, colWidths=[25 * mm, W - 25 * mm], hAlign="LEFT")
            meta_tbl.setStyle(TableStyle([
                ("BACKGROUND",    (0, 0), (0, -1), FLGREY),
                ("GRID",          (0, 0), (-1, -1), 0.25, LGREY),
                ("LEFTPADDING",   (0, 0), (-1, -1), 6),
                ("RIGHTPADDING",  (0, 0), (-1, -1), 6),
                ("TOPPADDING",    (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                ("VALIGN",        (0, 0), (-1, -1), "TOP"),
            ]))
            block.append(meta_tbl)
            block.append(Spacer(1, 4))

            # ── Steps table ───────────────────────────────────
            if exec_steps:
                block.append(Paragraph("Steps:", S["small"]))

                _th = _style("th", fontSize=8, fontName="Helvetica-Bold", textColor=DARK)
                trace_rows = [[
                    Paragraph("#",       _th),
                    Paragraph("Type",    _th),
                    Paragraph("Content", _th),
                    Paragraph("Status",  _th),
                ]]
                alt_style = []

                for row_i, s in enumerate(exec_steps, start=1):
                    stype   = (s.get("type") or "act").lower()
                    sstatus = (s.get("status") or "").lower()
                    is_failed = sstatus in ("failed", "error") or stype == "error"

                    type_label = stype.upper()
                    main_content = (s.get("content") or s.get("action") or "")
                    step_err = (
                        s.get("error") or s.get("error_message") or
                        s.get("error_detail") or ""
                    )

                    content_text = main_content[:300]
                    if is_failed and step_err:
                        content_text += f"\nError: {step_err[:200]}"
                    content_cell = Paragraph(
                        content_text,
                        _style(f"cc_{row_i}", fontSize=7, textColor=DARK, leading=10),
                    )

                    if sstatus in ("success", "passed"):
                        result_char  = "✓"
                        result_color = GREY
                    elif is_failed:
                        result_char  = "✗"
                        result_color = RED
                    else:
                        result_char  = "—"
                        result_color = GREY

                    trace_rows.append([
                        Paragraph(
                            str(s.get("order", row_i - 1)),
                            _style(f"tn_{row_i}", fontSize=7, textColor=GREY),
                        ),
                        Paragraph(
                            type_label,
                            _style(f"tt_{row_i}", fontSize=7, textColor=GREY),
                        ),
                        content_cell,
                        Paragraph(
                            result_char,
                            _style(f"ri_{row_i}", fontSize=9, textColor=result_color,
                                   alignment=TA_CENTER),
                        ),
                    ])

                    if row_i % 2 == 0:
                        alt_style.append(("BACKGROUND", (0, row_i), (-1, row_i), FLGREY))

                trace_tbl = Table(
                    trace_rows,
                    colWidths=[10 * mm, 13 * mm, 135 * mm, 12 * mm],
                    hAlign="LEFT",
                )
                trace_tbl.setStyle(TableStyle([
                    ("FONTSIZE",      (0, 0), (-1, -1), 7),
                    ("LEFTPADDING",   (0, 0), (-1, -1), 4),
                    ("RIGHTPADDING",  (0, 0), (-1, -1), 4),
                    ("TOPPADDING",    (0, 0), (-1, -1), 2),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
                    ("VALIGN",        (0, 0), (-1, -1), "TOP"),
                    ("GRID",          (0, 0), (-1, -1), 0.2, LGREY),
                    ("ALIGN",         (3, 0), (3, -1),  "CENTER"),
                    ("BACKGROUND",    (0, 0), (-1, 0),  FLGREY),
                ] + alt_style))
                block.append(trace_tbl)

            block.append(Spacer(1, 10))

            story.append(KeepTogether(block[:4]))
            story.extend(block[4:])

        # ── Footer ────────────────────────────────────────────
        story.append(Spacer(1, 12))
        story.append(_hr(LGREY, 0.5, 4))
        story.append(Paragraph(
            "Generated by TestForge AI — Intelligent Test Automation",
            S["footer"],
        ))

        doc.build(story)
        return buffer.getvalue()
